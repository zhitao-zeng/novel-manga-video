from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from .models import NovelDocument
from .splitting import split_episodes


SUPPORTED_SUFFIXES = {".txt", ".md", ".markdown", ".docx", ".pdf"}


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"cannot decode text file: {path}")


def _read_docx(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append("\t".join(cells))
    return "\n\n".join(blocks)


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(page for page in pages if page)
    if not text:
        raise ValueError("PDF contains no extractable text; OCR input is not supported in this build")
    return text


def read_novel(path: str | Path, novel_id: str, title: str | None = None) -> NovelDocument:
    source = Path(path).resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    suffix = source.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"unsupported file type {suffix}; supported: {sorted(SUPPORTED_SUFFIXES)}")
    if suffix in {".txt", ".md", ".markdown"}:
        text = _read_text(source)
    elif suffix == ".docx":
        text = _read_docx(source)
    else:
        text = _read_pdf(source)

    text = re.sub(r"\n{4,}", "\n\n", text).strip()
    episodes, chaptered = split_episodes(text)
    return NovelDocument(
        novel_id=str(novel_id),
        title=(title or source.stem).strip(),
        source_path=source,
        text=text,
        episodes=episodes,
        chaptered=chaptered,
    )
